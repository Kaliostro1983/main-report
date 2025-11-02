# src/reportgen/export/word_report.py
from __future__ import annotations

# from collections import OrderedDict, Counter
from pathlib import Path
from datetime import datetime
import re
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.shared import Pt

from src.armorkit.dates import combine_date_time, parse_period_from_filename, format_for_filename
from src.armorkit.domain.freqnorm import freq4_str
from src.armorkit.domain.callsigns import build_callsign_str_for_freq
from src.armorkit.domain.reference import (
    get_network_name_by_freq,
    full_tag_for_group,
    read_reference_sheet,
)

from src.armorkit.domain.schema import message_columns, COL_FREQ, COL_DATE, COL_TIME, COL_WHO, COL_TO
from src.armorkit.docxutils.images import insert_bearing_image
from src.armorkit.docxutils.styles import set_base_styles, add_title
from src.armorkit.docxutils.tables import set_col_widths, center_cell, vcenter, set_row_min_height
from src.armorkit.domain.intercepts import network_is_empty
from src.armorkit.docxutils.anchors import add_internal_link, bookmark

from src.armorkit.data_loader import load_inputs
from src.armorkit.normalize_freq import normalize_frequency_column, FREQ_NOT_FOUND
from src.reportgen.grouping import (
    unique_frequencies_with_counts,
    group_frequencies_by_tag,
)
from src.armorkit.settings import load_config

from src.armorkit.docxutils.safe_save import safe_save_docx

import logging

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)



def _render_frequency_section(doc: Document, freq4: str, count: int, li, cfg) -> None:
    # Якір
    title_p = doc.add_paragraph()
    anchor = f"freq-{freq4.replace('.', '_')}"
    bookmark(title_p, anchor)

    # Заголовок
    net_name = get_network_name_by_freq(freq4, li.reference_df)
    run = title_p.add_run(f"[{freq4}] - {net_name} - ({count})")
    run.bold = True
    run.font.size = Pt(12)

    # Еталонка
    ref_sheet = read_reference_sheet(freq4, li.freq_path)
    purpose = ref_sheet.get("Призначення") or "—"

    # Вузли: з головного листа; якщо порожньо — зі "Склад кореспондентів" еталонки
    nodes = "—"
    if "Частота" in li.reference_df.columns:
        df_ref = li.reference_df.copy()
        df_ref["__f4"] = df_ref["Частота"].map(freq4_str)
        m = df_ref[df_ref["__f4"] == freq4]
        if not m.empty:
            for col in ["Вузли зв’язку", "Вузли зв'язку", "Вузли", "Вузли звʹязку"]:
                if col in m.columns and pd.notna(m.iloc[0][col]) and str(m.iloc[0][col]).strip():
                    nodes = str(m.iloc[0][col]).strip()
                    break
    if nodes == "—":
        nodes = ref_sheet.get("Склад кореспондентів") or "—"

    # ТРИ абзаци з готовими значеннями
    doc.add_paragraph(f"Призначення радіомережі: {purpose}")
    doc.add_paragraph(f"Вузли зв’язку: {nodes}")
    
    # Позивні (беремо aliases з cfg.callsign_aliases)
    calls_text = build_callsign_str_for_freq(li.intercepts_df, freq4)
    doc.add_paragraph(f"Список позивних: {calls_text}")


            
# --- ПУБЛІЧНИЙ API: згенерувати DOCX-чернетку ---
__all__ = ["build_draft_docx"]

def create_move_report(config_path: str = "config.yml") -> str:

    """
    Генерує DOCX:
      - перша сторінка (огляд з клікабельними частотами)
      - розділи по кожній частоті (page break між ними)
    Повертає абсолютний шлях до збереженого файлу.
    """
    cfg = load_config(config_path)
    li = load_inputs(config_path)

    # нормалізуємо «Частота» в перехопленнях
    normalize_frequency_column(li.intercepts_df, li.reference_df)

    # групи та лічильники
    freqs, counts = unique_frequencies_with_counts(li.intercepts_df)
    allowed = (cfg.grouping or {}).get("allowed_tags", [])
    other   = (cfg.grouping or {}).get("other_bucket", "Інші радіомережі")
    groups = group_frequencies_by_tag(freqs, li.reference_df, allowed, other, cfg.grouping)

    # період з назви файла репорту
    period_start, period_end = parse_period_from_filename(li.report_path)

    # створюємо документ
    doc = Document()


    # 2) Детальні розділи по частотах
  # 1) Зібрати частоти, які МАЮТЬ коментовані перехоплення (зберігаємо порядок груп)
    pub_freqs: list[tuple[str, str]] = []  # (short_tag, freq4)
    for short_tag, flist in groups.items():
        for f in flist:
            # if not network_is_empty(li.intercepts_df, f):
            pub_freqs.append((short_tag, f))
            # else:
                # log.info("Секцію %s пропущено — немає перехоплень із коментарями.", f)

    # 2) Рендер секцій з розривом сторінки МІЖ ними
    for idx, (short_tag, f) in enumerate(pub_freqs, start=1):
        _render_frequency_section(doc, f, counts.get(f, 0), li, cfg)
        if idx < len(pub_freqs):
            doc.add_paragraph()
                
    # ...
    start_s = format_for_filename(period_start)
    end_s   = format_for_filename(period_end)
    file_name = f"Звіт з переміщення противника ({start_s} - {end_s}).docx"


    out_dir = Path(getattr(cfg.paths, "output_dir", "build"))
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / file_name

    saved = safe_save_docx(doc, out_path)  # як і раніше використовуємо безпечне збереження
    return str(saved.resolve())



