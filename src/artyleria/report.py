# -*- coding: utf-8 -*-
from __future__ import annotations
from pathlib import Path
from typing import Iterable, Mapping
from datetime import datetime
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TITLE = "Звіт з артилерії"

def _h1(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _h2(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)

def _p(doc: Document, text: str):
    doc.add_paragraph(text)

def _clean_header(body: str) -> str:
    """Прибрати службовий верхній рядок (ISO-datetime) та порожні рядки зверху."""
    if not body:
        return ""
    lines = [ln.rstrip() for ln in body.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and re.match(r"^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}(:\d{2})?", lines[0].strip()):
        lines.pop(0)
        while lines and not lines[0].strip():
            lines.pop(0)
    return "\n".join(lines).strip()

def build_docx(groups: Iterable[Mapping], out_path: str | Path) -> Path:
    """
    groups: ітерабель словників:
      {
        "freq": "300.3490",
        "name": "УКХ р/м ...",
        "image": Path|None,
        "intercepts": [ {"Дата": "...", "Час": "...", "text": "..."} , ... ]
      }
    """
    out_path = Path(out_path)
    doc = Document()

    # Заголовок
    today = datetime.now().strftime("%d.%m.%Y")
    _h1(doc, f"{TITLE} {today}")

    groups = list(groups)
    any_block = False

    for idx, g in enumerate(groups):
        inters = g.get("intercepts") or []
        if not inters:
            continue
        any_block = True

        freq = g["freq"]
        name = g["name"]
        _h2(doc, f"{freq} - {name}")

        # Зображення або заглушка
        img = g.get("image")
        if img and Path(img).exists():
            try:
                doc.add_picture(str(img), width=Inches(6.0))
            except Exception as e:
                print(f"[WARN] Не вдалось вставити зображення {img}: {e}")
                _p(doc, "р/м знаходиться на пеленгації")
        else:
            _p(doc, "р/м знаходиться на пеленгації")

        # Перехоплення (хронологія)
        for row in inters:
            d = str(row.get("Дата", "")).strip()
            t = str(row.get("Час", "")).strip()
            body = _clean_header(str(row.get("text", "")).strip())

            if not (body or d or t):
                continue
                
            if body:
                _p(doc, body)

        # відступ і розрив сторінки між блоками
        doc.add_paragraph("")
        if idx < len(groups) - 1:
            doc.add_page_break()

    if not any_block:
        _p(doc, "Немає перехоплень для артилерійських радіомереж на вхідній вибірці.")

    doc.save(out_path)
    return out_path
