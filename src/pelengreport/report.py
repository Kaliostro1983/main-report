# src/pelengreport/report.py
# -*- coding: utf-8 -*-
from __future__ import annotations
from datetime import datetime
from pathlib import Path
from typing import Iterable, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

# -------------------- helpers (портовано зі старої версії) --------------------
def _add_header(doc: Document, total_rows: int) -> None:
    # базовий стиль
    normal = doc.styles["Normal"]
    normal.font.size = Pt(12)
    normal.font.name = "Times New Roman"

    # правий верхній кут
    p = doc.add_paragraph("Форма 1.2.15")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")

    # заголовок
    p = doc.add_paragraph("ДОНЕСЕННЯ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # рядок з датою «станом на 24:00 <сьогодні> року»
    date_str = datetime.now().strftime("%d.%m.%Y")
    p = doc.add_paragraph(
        f"за результатами функціонування тактичної радіопеленгаторної мережі у зоні "
        f"відповідальності станом на 24:00 {date_str} року"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")


def _set_cell(cell, text: str, bold: bool = False,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              valign=WD_ALIGN_VERTICAL.CENTER):
    cell.text = ""
    lines = (text or "").split("\n")
    if not lines:
        lines = [""]

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
            p.text = line
        else:
            p = cell.add_paragraph(line)
        p.alignment = align
        if p.runs:
            p.runs[0].bold = bold
    cell.vertical_alignment = valign


def _add_body(doc: Document, total_pelengs: int) -> None:
    # 1. Склад сил і засобів…
    p = doc.add_paragraph("1. Склад сил і засобів, які розгорнуті для визначення місцеположення джерел (об’єктів) розвідки.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tbl1 = doc.add_table(rows=1, cols=5)
    tbl1.style = "Table Grid"
    hdrs1 = [
        "№ з/п",
        "Військова частина (підрозділ)",
        "Район розташування, номер бойового посту",
        "Озброєння, військова техніка, яка залучена",
        "Хід виконання розвідувальних завдань",
    ]
    for j, h in enumerate(hdrs1):
        _set_cell(tbl1.rows[0].cells[j], h, bold=True)

    row = tbl1.add_row()
    merged = row.cells[2].merge(row.cells[3]).merge(row.cells[4])
    _set_cell(row.cells[2], "3 АК", bold=True)
    _set_cell(row.cells[0], ""); _set_cell(row.cells[1], "")

    row = tbl1.add_row()
    _set_cell(row.cells[0], "1.")
    _set_cell(row.cells[1], "А3719\n(63 омбр)")
    _set_cell(row.cells[2], "МІКОЛАЇВКА,\nБП №0000", align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell(row.cells[3], "“Пластун”")
    _set_cell(row.cells[4], "Відповідно до плану бойового застосування", align=WD_ALIGN_PARAGRAPH.LEFT)

    row = tbl1.add_row()
    _set_cell(row.cells[0], "2.")
    _set_cell(row.cells[1], "А3719\n(63 омбр)")
    _set_cell(row.cells[2], "МАЯКИ,\nБП №0001", align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell(row.cells[3], "“Пластун”")
    _set_cell(row.cells[4], "Відповідно до плану бойового застосування", align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph("")

    # 2. Зміни в стані засобів…
    p = doc.add_paragraph(
        "2. Зміни в стані засобів пеленгування (вихід з ладу, зміна технічних позицій, "
        "розгортання нових засобів, втрати та заходи щодо відновлення)."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph("В стані та положенні засобів пеленгації змін немає.")
    if p.runs: p.runs[0].italic = True

    # 3. Загальна кількість…
    p = doc.add_paragraph(
        "3. Загальна кількість викритих (підтверджених) районів, кількість отриманих пеленгів (напрямків)."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tbl2 = doc.add_table(rows=1, cols=6)
    tbl2.style = "Table Grid"
    hdrs2 = [
        "№ з/п",
        "Військова частина (підрозділ)",
        "Район розташування, номер бойового посту",
        "Озброєння, військова техніка, яка залучена",
        "Кількість отриманих пеленгів (напрямків)",
        "Примітка",
    ]
    for j, h in enumerate(hdrs2):
        _set_cell(tbl2.rows[0].cells[j], h, bold=True)

    row = tbl2.add_row()
    row.cells[2].merge(row.cells[3]).merge(row.cells[4])
    _set_cell(row.cells[2], "3 АК", bold=True)
    _set_cell(row.cells[0], ""); _set_cell(row.cells[1], ""); _set_cell(row.cells[5], "")

    row = tbl2.add_row()
    _set_cell(row.cells[0], "1.")
    _set_cell(row.cells[1], "А3719\n(63 омбр)")
    _set_cell(row.cells[2], "МІКОЛАЇВКА,\nБП №0000", align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell(row.cells[3], "“Пластун”")
    _set_cell(row.cells[4], str(total_pelengs))
    _set_cell(row.cells[5], "")

    row = tbl2.add_row()
    _set_cell(row.cells[0], "2.")
    _set_cell(row.cells[1], "А3719\n(63 омбр)")
    _set_cell(row.cells[2], "МАЯКИ,\nБП №0001", align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell(row.cells[3], "“Пластун”")
    _set_cell(row.cells[4], "0")
    _set_cell(row.cells[5], "")

    doc.add_paragraph("")
    p = doc.add_paragraph("4. Результати визначення місцеположень джерел (об’єктів) розвідки.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("")


def _add_table(doc: Document, rows: Iterable[Mapping[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Частота (МГц)"
    hdr[2].text = "Назва підрозділу"
    hdr[3].text = "Дата та час"
    hdr[4].text = "Координати"

    for i, rec in enumerate(rows, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = str(rec.get("freq_or_mask", ""))  # маска або частота — як у джерелі
        cells[2].text = str(rec.get("unit_desc", ""))
        cells[3].text = str(rec.get("dt", ""))
        cells[4].text = str(rec.get("mgrs", ""))

    # центруємо де потрібно
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j in (0, 1, 3, 4):
                for pr in cell.paragraphs:
                    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx(records: list[Mapping[str, str]], out_path: str | Path) -> Path:
    """
    Формує DOCX «як у попередній версії»: шапка, розділи 1–3 (як у твоєму старому звіті),
    «4. …» + таблиця з даними (1 координата = 1 рядок).
    """
    out_path = Path(out_path)
    doc = Document()

    total = len(records)  # для секції 3 (кількість пеленгів/напрямків)
    _add_header(doc, total_rows=total)
    _add_body(doc, total_pelengs=total)
    _add_table(doc, records)

    # низ документа (підпис) — як у старій версії
    doc.add_paragraph("")
    p = doc.add_paragraph("Командир взводу РЕР _______СТЕПУРА Андрій Іванович_________")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(out_path)
    return out_path
