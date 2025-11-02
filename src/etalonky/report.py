# -*- coding: utf-8 -*-
from __future__ import annotations
from pathlib import Path
from typing import Iterable, Mapping
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



TITLE2 = "ДОНЕСЕННЯ"
TITLE3 = "Форма № 1.5.3/ГУР"
TITLE = "'Еталонні описи джерел радіовипромінювання противника у зоні відповідальності 63 ОМБр'"
SUB_ITLE = "Для КХ та УКХ діапазонів радіохвиль (з урахуванням виявлених джерел маневреними групами РЕР частин РЕР у районах виконання завдань)"



def _h1(doc: Document, text: str, isBold: bool = True, fontSize: int = 14, alignCenter: bool = True, alignRight: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = isBold
    r.font.size = Pt(fontSize)
    if alignRight:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignCenter:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _h2(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)


def _p(doc: Document, text: str):
    doc.add_paragraph(text)


def build_docx(sections: Iterable[Mapping], out_path: str | Path) -> Path:
    """
    sections: [{"freq": "300.3490", "lines": ["1. ...", "2. ...", ...]}, ...]
    """
    out_path = Path(out_path)
    doc = Document()

    # Шапка
    today = datetime.now().strftime("%d.%m.%Y")
    _h1(doc, f"{TITLE3}", alignRight=True, isBold=False)
    _h1(doc, f"{TITLE2}")
    _h1(doc, f"{TITLE}", isBold=False)
    _p(doc, "")

    sections = list(sections)
    if not sections:
        _p(doc, "Немає частот для публікації (статус 'Спостерігається' порожній).")
        doc.save(out_path)
        return out_path

    # Кожна еталонка — з нової сторінки
    for idx, s in enumerate(sections):
        freq = s["freq"]
        lines = s.get("lines", [])

        # _h2(doc, f"[{freq}]")
        _h1(doc, f"{SUB_ITLE}", alignCenter=False, fontSize=12)

        for line in lines:
            _p(doc, line)

        if idx < len(sections) - 1:
            doc.add_page_break()

    doc.save(out_path)
    return out_path
