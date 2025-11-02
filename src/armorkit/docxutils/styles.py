import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



# -----------------------
# Базові стилі/утиліти
# -----------------------
def set_base_styles(doc: Document):
    style = doc.styles['Normal']
    style.font.size = Pt(12)

def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return p


def add_right_text(doc: Document, text: str, *, bold: bool = False, font_size_pt: int = 12):
    """
    Додає праворуч вирівняний рядок (наприклад 'Форма №1.2.17/ОСУВ')
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size_pt)
    return p

def add_center_block(doc: Document, lines: list[str], *, bold: bool = False, font_size_pt: int = 12):
    """
    Додає центрований блок з кількох рядків.
    Використовується для:
    - 'Донесення ... у зоні ... “Кремінна”'
    - '(станом на ...)'
    """
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(font_size_pt)
    return

def add_block_title(doc: Document, text: str):
    """
    Заголовок інформаційного блоку перед таблицею (зліва, без болду)
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = False
    run.font.size = Pt(12)
    return p

def add_footer_signature(doc: Document):
    """
    Службовий підпис наприкінці звіту
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(
        "Командир взводу РЕР ____________СТЕПУРА Андрій Іванович________________________"
    )
    run.font.size = Pt(12)
    return p
