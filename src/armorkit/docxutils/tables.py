import logging
from docx import Document   
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import pandas as pd


def set_col_widths(table, factors):
    total = sum(factors)
    total_inches = 6.5
    widths = [Inches(total_inches * f / total) for f in factors]
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = w

def center_cell(cell):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def vcenter(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_row_min_height(row, cm: float = 0.9):
    row.height = Cm(cm)
    row.height_rule = WD_ROW_HEIGHT.AT_LEAST
    

def _ensure_tblPr(tbl_element):
    """
    Повертає (або створює) вузол <w:tblPr> для таблиці.
    Це сумісно з різними версіями python-docx, де немає get_or_add_tblPr().
    """
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_element.insert(0, tblPr)
    return tblPr


def _set_table_borders(table):
    """
    Додає чорні межі (бордери) навколо і всередині таблиці:
    top, bottom, left, right, insideH, insideV.
    """
    tbl = table._element  # CT_Tbl

    tblPr = _ensure_tblPr(tbl)

    # створити/перезаписати w:tblBorders
    tblBorders = OxmlElement("w:tblBorders")

    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{border_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "8")        # товщина лінії (1/8 точки)
        edge.set(qn("w:color"), "000000")
        tblBorders.append(edge)

    # приєднати до tblPr
    # якщо в tblPr вже є borders, треба його прибрати
    for child in list(tblPr):
        if child.tag == qn("w:tblBorders"):
            tblPr.remove(child)
    tblPr.append(tblBorders)


def add_intercepts_table(doc, rows_data: list[tuple[str, str]]):
    """
    Створює таблицю з колонками:
      'р\\обмін' | 'примітки'
    Додає рамки, виставляє шрифти.
    rows_data — список кортежів (r_obmin, prymitky)
    """

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "р\\обмін"
    hdr_cells[1].text = "примітки"

    for r_obmin, note in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = "" if (pd.isna(r_obmin) or r_obmin is None) else str(r_obmin)
        row_cells[1].text = "" if (pd.isna(note) or note is None) else str(note)

    # оформлення тексту
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                # перший рядок (заголовок таблиці) центруємо, інше — ліворуч
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(12)

    # рамки навколо/всередині таблиці
    _set_table_borders(table)

    return table