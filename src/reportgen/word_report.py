from __future__ import annotations
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from pathlib import Path

def export_docx(df: pd.DataFrame, out_path: str, title: str = "Аналітичний звіт",
                subtitle: Optional[str] = None) -> str:
    doc = Document()
    # Заголовок
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Згенеровано: {datetime.now():%Y-%m-%d %H:%M:%S}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Короткі метрики
    doc.add_heading("Резюме", level=1)
    summary = {
        "Кількість рядків": len(df),
        "Кількість колонок": len(df.columns),
        "Джерел (файлів)": df["__source__"].nunique() if "__source__" in df.columns else "—",
    }
    for k, v in summary.items():
        doc.add_paragraph(f"{k}: {v}")

    # Таблиця (обрізаємо до перших 20 рядків для прикладу)
    doc.add_heading("Вибірка даних (top 20)", level=1)
    head_df = df.head(20).copy()
    table = doc.add_table(rows=1, cols=len(head_df.columns))
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(head_df.columns):
        hdr_cells[i].text = str(c)
    for _, row in head_df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(head_df.columns):
            cells[i].text = str(row.get(c, ""))

    # Збереження
    out_path = str(Path(out_path))
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
