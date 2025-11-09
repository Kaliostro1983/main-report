# src/reportgen/export/word_report.py
from __future__ import annotations

# from collections import OrderedDict, Counter
from pathlib import Path
from datetime import datetime
import re
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.shared import Pt

from src.armorkit.dates import combine_date_time, parse_period_from_filename, format_for_filename
from src.armorkit.domain.freqnorm import freq4_str
from src.armorkit.domain.callsigns import build_callsign_str_for_freq
from src.armorkit.domain.reference import (
    get_network_name_by_freq,
    full_tag_for_group,
    read_reference_sheet,
)

from src.armorkit.domain.schema import message_columns, COL_FREQ, COL_DATE, COL_TIME, COL_WHO, COL_TO
from src.armorkit.docxutils.images import insert_bearing_image
from src.armorkit.docxutils.styles import set_base_styles, add_title
from src.armorkit.docxutils.tables import set_col_widths, center_cell, vcenter, set_row_min_height
from src.armorkit.domain.intercepts import network_is_empty
from src.armorkit.docxutils.anchors import add_internal_link, bookmark

from src.armorkit.data_loader import load_inputs
from src.armorkit.normalize_freq import normalize_frequency_column, FREQ_NOT_FOUND
from src.reportgen.grouping import (
    unique_frequencies_with_counts,
    group_frequencies_by_tag,
)
from src.armorkit.settings import load_config

from src.armorkit.docxutils.safe_save import safe_save_docx

import logging

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)


def _append_executor_block(doc: Document) -> None:
    """Додає примітку і рядок виконавця в кінці звіту (без розриву сторінки)."""
    note = (
        "* Добування розвідувальної інформації про противника здійснюється підрозділами "
        "РЕР, розгорнутими в смугах відповідальності ОТУ та надається для первинної обробки "
        "у відповідні підрозділи РЕР центрів (відділів) розвідки ОТУ, де інформація "
        "узагальнюється та надається короткий опис подій.\n\n"
        "Матеріали радіоперехоплень в узагальненому вигляді від підрозділів РЕР ЦР (ВР) ОТУ у "
        "визначений час надаються черговому розвідки ОКП ОСУВ “Хортиця”. Інформація, яка "
        "потребує невідкладного доведення до начальника центру розвідки ОКП ОСУВ "
        "“Хортиця”, надається черговому розвідки ОКП ОСУВ “Хортиця” негайно по тлф з "
        "подальшим документальним підтвердженням."
    )
    p1 = doc.add_paragraph(note)
    for r in p1.runs:
        r.italic = True
        r.font.size = Pt(12)

    doc.add_paragraph()  # порожній рядок

    p2 = doc.add_paragraph(
        "Командир взводу РЕР _________СТЕПУРА Андрій Іванович__________________"
    )
    for r in p2.runs:
        r.font.size = Pt(12)



def _render_frequency_section(doc: Document, freq4: str, count: int, li, cfg) -> None:
    # Якір
    title_p = doc.add_paragraph()
    anchor = f"freq-{freq4.replace('.', '_')}"
    bookmark(title_p, anchor)

    # Заголовок
    net_name = get_network_name_by_freq(freq4, li.reference_df)
    run = title_p.add_run(f"[{freq4}] - {net_name} - ({count})")
    run.bold = True
    run.font.size = Pt(12)

    # Еталонка
    ref_sheet = read_reference_sheet(freq4, li.freq_path)
    purpose = ref_sheet.get("Призначення") or "—"

    # Вузли: з головного листа; якщо порожньо — зі "Склад кореспондентів" еталонки
    nodes = "—"
    if "Частота" in li.reference_df.columns:
        df_ref = li.reference_df.copy()
        df_ref["__f4"] = df_ref["Частота"].map(freq4_str)
        m = df_ref[df_ref["__f4"] == freq4]
        if not m.empty:
            for col in ["Вузли зв’язку", "Вузли зв'язку", "Вузли", "Вузли звʹязку"]:
                if col in m.columns and pd.notna(m.iloc[0][col]) and str(m.iloc[0][col]).strip():
                    nodes = str(m.iloc[0][col]).strip()
                    break
    if nodes == "—":
        nodes = ref_sheet.get("Склад кореспондентів") or "—"

    # ТРИ абзаци з готовими значеннями
    doc.add_paragraph(f"Призначення радіомережі: {purpose}")
    doc.add_paragraph(f"Вузли зв’язку: {nodes}")
    
    # Позивні (беремо aliases з cfg.callsign_aliases)
    calls_text = build_callsign_str_for_freq(li.intercepts_df, freq4)
    doc.add_paragraph(f"Список позивних: {calls_text}")

    # Далі — як було: таблиця з 2 колонок тільки для перехоплень з коментарем
    doc.add_paragraph("Найважливіші перехоплення з коментарями:").runs[0].bold = True

    msg_col, cmt_col = message_columns(li.intercepts_df)
    df = li.intercepts_df.copy()
    df["__f4"] = df["Частота"].map(freq4_str)
    part = df[df["__f4"] == freq4].copy()

    if cmt_col:
        cm = part[cmt_col].astype(str).fillna("").str.strip().replace({"nan": "", "None": "", "NONE": ""})
        part = part[cm.ne("")]
    else:
        part = part.iloc[0:0]

    if part.empty:
        # якщо з якихось причин сюди дійшли без записів — просто не друкуємо пусту таблицю
        log.info("Секція %s: відсутні перехоплення з коментарями (таблиця пропущена).", freq4)
        # doc.add_page_break()
        return

    if all(c in part.columns for c in ("Дата", "Час")):
        part["__dt"] = [combine_date_time(d, t) for d, t in zip(part["Дата"], part["Час"])]
        part = part.sort_values("__dt", kind="stable")

    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Перехоплення"; hdr[1].text = "Коментар"
    for c in hdr: center_cell(c); vcenter(c)
    set_row_min_height(t.rows[0], cm=0.9); set_col_widths(t, [3.6, 2.4])

    for _, row in part.iterrows():
        msg = str(row[msg_col]).strip() if msg_col and pd.notna(row.get(msg_col)) else ""
        cmt = str(row[cmt_col]).strip() if cmt_col and pd.notna(row.get(cmt_col)) else ""
        cells = t.add_row().cells
        cells[0].text = msg
        cells[1].text = cmt
        set_row_min_height(t.rows[-1], cm=0.9)

    insert_bearing_image(doc, freq4)
    # doc.add_page_break()





# -----------------------
# Перша сторінка (ОГЛЯД) — БЕЗ ЗМІН ВІД ТВОЄЇ ОСТАННЬОЇ ВЕРСІЇ,
# але замість plain-text частоти додаємо КЛІКАЛЬНЕ посилання.
# -----------------------
def _render_overview_page(doc: Document, cfg, li, groups, counts, period_start, period_end):
    set_base_styles(doc)

    add_title(doc, "Донесення")

    psub = doc.add_paragraph("за результатами ведення радіоелектронної розвідки\n"
                             "у зоні відповідальності тактичної групи “Кремінна”")
    psub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in psub.runs:
        r.bold = True; r.font.size = Pt(12)

    pper = doc.add_paragraph()
    pper.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pper.add_run(f"(з {period_start.split(' ')[1]} {period_start.split(' ')[0]} "
                       f"по {period_end.split(' ')[1]} {period_end.split(' ')[0]} року)")
    run.font.size = Pt(12)

    doc.add_paragraph()
    total_intercepts = int(len(li.intercepts_df))
    pinfo = doc.add_paragraph()
    pinfo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pinfo.add_run(f"За поточний період отримано {total_intercepts} перехоплень.")
    rr.bold = True; rr.font.size = Pt(12)

    doc.add_paragraph()
    title_tbl = doc.add_paragraph("Активність радіомереж:")
    title_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_tbl.runs[0].bold = True; title_tbl.runs[0].font.size = Pt(12)
    doc.add_paragraph()

       # Таблиця: № | Частота | Радіомережа | Перехоплення
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Частота"
    hdr[2].text = "Радіомережа"
    hdr[3].text = "Перехоплення"

    # заголовки по центру (горизонтально і вертикально) + висота
    for c in hdr:
        center_cell(c)
        vcenter(c)
    set_row_min_height(t.rows[0], cm=0.9)

    # Ширини колонок:
    # 1 — в 4 рази менша; 2 — на 35% менша; 3 — вдвічі більша; 4 — в 4 рази менша.
    # Пропорції: [0.25, 0.65, 2.0, 0.25]
    set_col_widths(t, [0.25, 0.65, 2.0, 0.25])

    row_counter = 1
    for short_tag, flist in groups.items():
        # повний напис «Хто» (за довідником)
        full_tag = full_tag_for_group(flist, li.reference_df, short_tag)

        # рядок-заголовок групи (злиті комірки, жирним, по центру)
        r = t.add_row()
        c = r.cells
        c[0].merge(c[1]).merge(c[2]).merge(c[3])
        c[0].text = f"Радіомережі {full_tag}"
        # стилізація комірки заголовка групи
        p = c[0].paragraphs[0]
        if p.runs:
            p.runs[0].bold = True
        center_cell(c[0])
        vcenter(c[0])
        set_row_min_height(t.rows[-1], cm=0.9)

        if not flist:
            # порожня група
            r = t.add_row().cells
            r[2].text = "Не виявлено"
            if r[2].paragraphs[0].runs:
                r[2].paragraphs[0].runs[0].italic = True
            # центруємо 1,2,4 колонки
            for idx in (0, 1, 3):
                center_cell(r[idx])
                vcenter(r[idx])
            set_row_min_height(t.rows[-1], cm=0.9)
            continue

        # рядки з даними по частотах
        for f in flist:
            r = t.add_row().cells
            # №
            r[0].text = str(row_counter)
            center_cell(r[0]); vcenter(r[0])

            # Частота як клікабельний лінк на закладку розділу частоти
            if not network_is_empty(li.intercepts_df, f):
                anchor = f"freq-{f.replace('.', '_')}"
                add_internal_link(r[1], f, anchor)
            else:
                r[1].text = f  # просто текст без гіперпосилання
            center_cell(r[1]); vcenter(r[1])

            # Назва мережі (з довідника)
            r[2].text = get_network_name_by_freq(f, li.reference_df)

            # Кількість перехоплень
            r[3].text = str(counts.get(f, 0))
            center_cell(r[3]); vcenter(r[3])

            set_row_min_height(t.rows[-1], cm=0.9)
            row_counter += 1
            
    

            
# --- ПУБЛІЧНИЙ API: згенерувати DOCX-чернетку ---
__all__ = ["build_draft_docx"]

def build_draft_docx(config_path: str = "config.yml") -> str:

    """
    Генерує DOCX:
      - перша сторінка (огляд з клікабельними частотами)
      - розділи по кожній частоті (page break між ними)
    Повертає абсолютний шлях до збереженого файлу.
    """
    cfg = load_config(config_path)
    li = load_inputs(config_path)

    # нормалізуємо «Частота» в перехопленнях
    normalize_frequency_column(li.intercepts_df, li.reference_df, li.masks_df)

    # групи та лічильники
    freqs, counts = unique_frequencies_with_counts(li.intercepts_df)
    allowed = (cfg.grouping or {}).get("allowed_tags", [])
    other   = (cfg.grouping or {}).get("other_bucket", "Інші радіомережі")
    groups = group_frequencies_by_tag(freqs, li.reference_df, allowed, other, cfg.grouping)

    # період з назви файла репорту
    period_start, period_end = parse_period_from_filename(li.report_path)

    # створюємо документ
    doc = Document()

    # 1) Перша сторінка-огляд
    _render_overview_page(doc, cfg, li, groups, counts, period_start, period_end)
    doc.add_page_break()

    # 2) Детальні розділи по частотах
  # 1) Зібрати частоти, які МАЮТЬ коментовані перехоплення (зберігаємо порядок груп)
    pub_freqs: list[tuple[str, str]] = []  # (short_tag, freq4)
    for short_tag, flist in groups.items():
        for f in flist:
            if not network_is_empty(li.intercepts_df, f):
                pub_freqs.append((short_tag, f))
            else:
                log.info("Секцію %s пропущено — немає перехоплень із коментарями.", f)

    # 2) Рендер секцій з розривом сторінки МІЖ ними
    for idx, (short_tag, f) in enumerate(pub_freqs, start=1):
        _render_frequency_section(doc, f, counts.get(f, 0), li, cfg)
        if idx < len(pub_freqs):
            doc.add_page_break()
                
    # після останньої секції — примітка + підпис (без page break)
    _append_executor_block(doc)

    # ...
    start_s = format_for_filename(period_start)
    end_s   = format_for_filename(period_end)
    file_name = f"Звіт РЕР ({start_s} - {end_s}).docx"


    out_dir = Path(getattr(cfg.paths, "output_dir", "build"))
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / file_name

    saved = safe_save_docx(doc, out_path)  # як і раніше використовуємо безпечне збереження
    return str(saved.resolve())



