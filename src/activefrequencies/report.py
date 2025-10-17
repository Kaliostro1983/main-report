# src/reportgen/export/word_report.py
from __future__ import annotations

# from collections import OrderedDict, Counter
from pathlib import Path
from datetime import datetime
import re
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.armorkit.dates import parse_period_from_filename, format_for_filename
from src.armorkit.domain.reference import (
    get_network_name_by_freq,
    full_tag_for_group,
)
from src.armorkit.domain.schema import message_columns, COL_FREQ, COL_DATE, COL_TIME, COL_WHO, COL_TO
from src.armorkit.docxutils.styles import set_base_styles, add_title
from src.armorkit.docxutils.tables import set_col_widths, center_cell, vcenter, set_row_min_height
from src.armorkit.domain.intercepts import network_is_empty


from src.armorkit.data_loader import load_inputs
from src.armorkit.normalize_freq import normalize_frequency_column, FREQ_NOT_FOUND
from src.reportgen.grouping import (
    unique_frequencies_with_counts,
    group_frequencies_by_tag,
)
from src.reportgen.settings import load_config

from src.armorkit.docxutils.safe_save import safe_save_docx

import logging
log = logging.getLogger(__name__)



def _append_executor_block(doc: Document) -> None:
    """Додає примітку і рядок виконавця в кінці звіту (без розриву сторінки)."""
    

    doc.add_paragraph()  # порожній рядок

    p2 = doc.add_paragraph(
        "Командир взводу РЕР _________СТЕПУРА Андрій Іванович__________________"
    )
    for r in p2.runs:
        r.font.size = Pt(12)



def _add_internal_link(cell, text: str, anchor: str):
    """
    Додає внутрішній лінк на закладку 'anchor' всередині осередка таблиці.
    """
    p = cell.paragraphs[0]
    # очистимо існуючий текст
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor)
    # стиль гіперлінка без синього/підкреслення: створимо run і налаштуємо вручну
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # прибираємо підкреслення й колір
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'none')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '000000')
    rPr.append(u); rPr.append(color)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(rPr); new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)

# -----------------------
# Перша сторінка (ОГЛЯД) — БЕЗ ЗМІН ВІД ТВОЄЇ ОСТАННЬОЇ ВЕРСІЇ,
# але замість plain-text частоти додаємо КЛІКАЛЬНЕ посилання.
# -----------------------
def _render_overview_page(doc: Document, cfg, li, groups, counts, period_start, period_end):
    set_base_styles(doc)

    add_title(doc, "Активні мережі (63 омсбр)")
    add_title(doc, f"станом на {period_end.split(' ')[0]}")

    psub = doc.add_paragraph("за результатами ведення радіоелектронної розвідки\n"
                             "у зоні відповідальності тактичної групи “Кремінна”")
    psub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in psub.runs:
        r.bold = True; r.font.size = Pt(12)

    pper = doc.add_paragraph()
    pper.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pper.add_run(f"(В пероід з {period_start.split(' ')[1]} {period_start.split(' ')[0]} "
                       f"по {period_end.split(' ')[1]} {period_end.split(' ')[0]} виявлено функціонування наступних радіомереж)")
    run.font.size = Pt(12)

    doc.add_paragraph()
 

       # Таблиця: № | Частота | Радіомережа | Перехоплення
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Частота"
    hdr[2].text = "Радіомережа"


    # заголовки по центру (горизонтально і вертикально) + висота
    for c in hdr:
        center_cell(c)
        vcenter(c)
    set_row_min_height(t.rows[0], cm=0.9)

    # Ширини колонок:
    # 1 — в 4 рази менша; 2 — на 35% менша; 3 — вдвічі більша; 4 — в 4 рази менша.
    # Пропорції: [0.25, 0.65, 2.0, 0.25]
    set_col_widths(t, [0.1, 0.45, 4.0])

    row_counter = 1
    for short_tag, flist in groups.items():
        # повний напис «Хто» (за довідником)
        full_tag = full_tag_for_group(flist, li.reference_df, short_tag)

        # рядок-заголовок групи (злиті комірки, жирним, по центру)
        r = t.add_row()
        c = r.cells
        c[0].merge(c[1]).merge(c[2])
        c[0].text = f"Радіомережі {full_tag}"
        # стилізація комірки заголовка групи
        p = c[0].paragraphs[0]
        if p.runs:
            p.runs[0].bold = True
        center_cell(c[0])
        vcenter(c[0])
        set_row_min_height(t.rows[-1], cm=0.9)

        if not flist:
            # порожня група
            r = t.add_row().cells
            r[2].text = "Не виявлено"
            if r[2].paragraphs[0].runs:
                r[2].paragraphs[0].runs[0].italic = True
            # центруємо 1,2,4 колонки
            for idx in (0, 1, 3):
                center_cell(r[idx])
                vcenter(r[idx])
            set_row_min_height(t.rows[-1], cm=0.9)
            continue

        # рядки з даними по частотах
        for f in flist:
            r = t.add_row().cells
            # №
            r[0].text = str(row_counter)
            center_cell(r[0]); vcenter(r[0])

            # Частота як клікабельний лінк на закладку розділу частоти
            if not network_is_empty(li.intercepts_df, f):
                anchor = f"freq-{f.replace('.', '_')}"
                _add_internal_link(r[1], f, anchor)
            else:
                r[1].text = f  # просто текст без гіперпосилання
            center_cell(r[1]); vcenter(r[1])

            # Назва мережі (з довідника)
            r[2].text = get_network_name_by_freq(f, li.reference_df)

            set_row_min_height(t.rows[-1], cm=0.9)
            row_counter += 1
            

            
# --- ПУБЛІЧНИЙ API: згенерувати DOCX-чернетку ---
__all__ = ["build_active_frequencies_docx"]

def build_active_frequencies_docx(config_path: str = "config.yml") -> str:
    
    import logging
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    """
    Генерує DOCX:
      - перша сторінка (огляд з клікабельними частотами)
      - розділи по кожній частоті (page break між ними)
    Повертає абсолютний шлях до збереженого файлу.
    """
    cfg = load_config(config_path)
    li = load_inputs(config_path)

    # нормалізуємо «Частота» в перехопленнях
    normalize_frequency_column(li.intercepts_df, li.reference_df)

    # групи та лічильники
    freqs, counts = unique_frequencies_with_counts(li.intercepts_df)
    allowed = (cfg.grouping or {}).get("allowed_tags", [])
    other   = (cfg.grouping or {}).get("other_bucket", "Інші радіомережі")
    groups = group_frequencies_by_tag(freqs, li.reference_df, allowed, other, cfg.grouping)

    # період з назви файла репорту
    period_start, period_end = parse_period_from_filename(li.report_path)

    # створюємо документ
    doc = Document()

    # 1) Перша сторінка-огляд
    _render_overview_page(doc, cfg, li, groups, counts, period_start, period_end)
    _append_executor_block(doc)

    # збереження

    end_s   = format_for_filename(period_end)
    file_name = f"Активні мережі (63 омсбр) {end_s[:8]}.docx"

    out_dir = Path(getattr(cfg.paths, "output_dir", "build"))
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / file_name

    saved = safe_save_docx(doc, out_path)  # як і раніше використовуємо безпечне збереження
    return str(saved.resolve())



