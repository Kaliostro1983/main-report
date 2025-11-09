# -*- coding: utf-8 -*-
"""
ENEMIES v3 — генерація DOCX-звіту про ворожі радіомережі.
"""
from __future__ import annotations
from pathlib import Path
from datetime import datetime
from collections import OrderedDict
import logging
import pandas as pd

log = logging.getLogger("enemies_report")
if not log.handlers:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log.info("=== ENEMIES v3: запуск модуля ===")

# ---- імпорти проєкту ----
from src.armorkit.data_loader import load_inputs
from src.armorkit.normalize_freq import is_real_freq, get_true_freq_by_mask, normalize_frequency_column
from src.armorkit.domain.freqnorm import freq4_str
from src.armorkit.domain.reference import get_network_name_by_freq
from src.armorkit.docxutils.safe_save import safe_save_docx

# DOCX
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.armorkit.docxutils.tables import center_cell, vcenter, set_col_widths, set_row_min_height


# ---------- допоміжні ----------
def read_freq_tokens(freq_file: Path) -> list[str]:
    text = freq_file.read_text(encoding="utf-8", errors="ignore")
    toks = [t.strip() for t in text.replace("\n", ",").split(",")]
    return [t for t in toks if t]


def tokens_to_freq4(tokens: list[str], reference_df: pd.DataFrame) -> list[str]:
    out = OrderedDict()
    for tok in tokens:
        true_f = tok if is_real_freq(tok) else get_true_freq_by_mask(tok, reference_df)
        out[freq4_str(true_f)] = None
    return list(out.keys())


def activity_period_for_freq4(intercepts_df: pd.DataFrame, f4: str) -> tuple[str, str]:
    """
    Повертає (перший час, останній час) для частоти f4 у форматі HH:MM.
    Якщо перехоплень немає — ('-', '-').
    Очікує колонки: 'Частота', 'Дата', 'Час'.
    """
    need = {"Частота", "Дата", "Час"}
    if not need.issubset(intercepts_df.columns):
        return "-", "-"

    df = intercepts_df.copy()
    df["__f4"] = df["Частота"].map(freq4_str)
    part = df.loc[df["__f4"] == f4].copy()
    if part.empty:
        return "-", "-"

    # нормалізуємо дату/час і збираємо єдиний datetime
    part["__date"] = pd.to_datetime(part["Дата"], dayfirst=True, errors="coerce")
    part["__time"] = part["Час"].astype(str).str.strip()
    part.loc[:, "__dt"] = pd.to_datetime(
        part["__date"].dt.strftime("%Y-%m-%d") + " " + part["__time"], errors="coerce"
    )

    part = part.loc[part["__dt"].notna()].sort_values("__dt", kind="stable")
    if part.empty:
        return "-", "-"

    # ЛИШЕ ЧАСИ
    tfmt = "%H:%M"
    start_t = part.iloc[0]["__dt"].strftime(tfmt)
    end_t   = part.iloc[-1]["__dt"].strftime(tfmt)
    return start_t, end_t



def mask3_from_reference(reference_df: pd.DataFrame, f4: str) -> str | None:
    df = reference_df.copy()
    if "Частота" not in df.columns:
        return None
    df["__f4"] = df["Частота"].map(freq4_str)
    m = df[df["__f4"] == f4]
    if m.empty:
        return None
    def _norm(val):
        s = str(val).strip()
        try: return f"{float(s.replace(',', '.')):.3f}"
        except Exception: return s
    for col in ("Маска_3", "Маска_Ш"):
        if col in m.columns and pd.notna(m.iloc[0][col]) and str(m.iloc[0][col]).strip():
            return _norm(m.iloc[0][col])
    return None


# ---------- ЕТАЛОНКИ: читаємо Категорія/Значення прямо з аркуша ----------
def read_ref_fields(freq_book_path: Path, f4: str) -> tuple[str, str, str, str, str, str]:
    """
    Повертає: (modulation, nature, main_vz, sub_vz, area, period)
    Підтримує:
      A) нормальні колонки з цими назвами
      Б) таблицю «№ | Категорія | Значення» (порядок колонок довільний)
    """
    try:
        xls = pd.ExcelFile(freq_book_path, engine="openpyxl")
    except Exception as e:
        log.warning("Довідник не відкрито: %s", e)
        return "—", "—", "—", "—", "—", "—"

    # 1) знайдемо аркуш: точний або «м’які» варіанти (145.9500 -> 145.95 тощо)
    sheet_names = [str(s).strip() for s in xls.sheet_names]
    target = None
    if f4 in sheet_names:
        target = f4
    else:
        for v in {f4, f"{float(f4):.4f}", f"{float(f4):.3f}", f"{float(f4):.2f}", str(float(f4))}:
            if v in sheet_names:
                target = v
                break
    if target is None:
        log.info("Еталонний аркуш відсутній: %s", f4)
        return "—", "—", "—", "—", "—", "—"

    try:
        df = pd.read_excel(xls, sheet_name=target, dtype=str)
    except Exception as e:
        log.warning("Аркуш '%s' не зчитано: %s", target, e)
        return "—", "—", "—", "—", "—", "—"

    # Нормалізація назв колонок
    cols_norm = {i: (str(c).strip() if c is not None else "") for i, c in enumerate(df.columns)}
    cols_lc = {i: cols_norm[i].lower() for i in cols_norm}

    # ---- ВАРІАНТ Б: "Категорія | Значення" у будь-яких позиціях ----
    def find_col(substrs: list[str]) -> int | None:
        for i, name_lc in cols_lc.items():
            if any(s in name_lc for s in substrs):
                return i
        return None

    cat_idx = find_col(["категор", "category"])
    val_idx = find_col(["значен", "value"])

    if cat_idx is not None and val_idx is not None:
        keys = (
            df.iloc[:, cat_idx]
              .astype(str)
              .str.strip()
              .str.replace(r"\s+", " ", regex=True)
              .str.lower()
        )
        vals = df.iloc[:, val_idx].apply(lambda x: "" if pd.isna(x) else str(x).strip())
        kv = {}
        for k, v in zip(keys, vals):
            if k and v:
                kv[k] = v.replace("\r\n", "\n").replace("\r", "\n")

        def pick_kv(*alts: str) -> str:
            for a in alts:
                a_lc = a.lower()
                if a_lc in kv and kv[a_lc]:
                    return kv[a_lc]
            return "—"

        modulation = pick_kv("вид передачі", "вид модуляції", "модуляція")
        nature     = pick_kv("характер роботи", "характер радіообміну", "призначення")
        main_vz    = pick_kv("склад кореспондентів", "вузли зв’язку")
        sub_vz     = pick_kv("позивні", "позивні (еталон)")
        area       = pick_kv("район функціонування/розгортання", "район функціонування", "район розгортання")
        period     = pick_kv("період функціонування", "період роботи", "період активності")

        return modulation, nature, main_vz, sub_vz, area, period

    # ---- ВАРІАНТ А: беремо напряму з колонок (якщо хтось зробив "плоску" таблицю) ----
    def col_first(*names: str) -> str:
        for n in names:
            if n in df.columns:
                ser = df[n].dropna()
                if not ser.empty:
                    s = str(ser.iloc[0]).strip().replace("\r\n", "\n").replace("\r", "\n")
                    if s:
                        return s
        return "—"

    modulation = col_first("Вид передачі", "Вид модуляції", "Модуляція")
    nature     = col_first("Характер роботи", "Характер радіообміну", "Призначення")
    main_vz    = col_first("Склад кореспондентів", "Вузли зв’язку")
    sub_vz     = col_first("Позивні", "Позивні (еталон)")
    area       = col_first("Район функціонування/розгортання", "Район функціонування", "Район розгортання")
    period     = col_first("Період функціонування", "Період роботи", "Період активності")

    return modulation, nature, main_vz, sub_vz, area, period



# ---------- DOCX ----------
def add_header(doc: Document, date_str: str, unit_name: str = "63 омсбр"):
    h1 = doc.add_paragraph(f"Ворожі радіомережі в тилу ({unit_name})")
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h1.runs[0]; r.bold = True; r.font.size = Pt(14)

    h2 = doc.add_paragraph(f"станом на {date_str}")
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.runs[0].font.size = Pt(13)

    p = doc.add_paragraph(
        "В результаті проведення радіорозвідки у смузі оборони "
        f"{unit_name}, виявлено функціонування ворожих радіомереж у тилових районах СОУ."
    )
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)

    
    doc.add_paragraph()


def add_overview_table(doc: Document, rows: list[dict]):
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "№"; hdr[1].text = "Частота (МГц)"; hdr[2].text = "Вид модуляції"
    hdr[3].text = "Період активності"; hdr[4].text = "Координати"; hdr[5].text = "Примітки"

    for i, row in enumerate(rows, start=1):
        c = t.add_row().cells
        c[0].text = str(i)
        c[1].text = row["freq4"]
        c[2].text = row.get("modulation", "—")
        c[3].text = "—" if (row.get("period_start") == "-" and row.get("period_end") == "-") \
                    else f"{row['period_start']} — {row['period_end']}"
        c[4].text = "—"; c[5].text = "—"

    set_col_widths(t, [0.8, 1.4, 1.5, 2.6, 1.4, 1.2])
    for row in t.rows:
        set_row_min_height(row, cm=0.9)
        for cell in row.cells:
            center_cell(cell); vcenter(cell)
    doc.add_paragraph()


def add_network_blocks(doc: Document, items: list[dict]):
    for it in items:
        head = f"[{it['freq4']}]"
        if it.get("mask3"): head += f" ({it['mask3']})"
        head += f" - {it.get('network_name','—')}"
        p = doc.add_paragraph(); r = p.add_run(head); r.bold = True; r.font.size = Pt(12)

        if it.get("area") and it["area"] != "—":
            doc.add_paragraph(f"Район функціонування: {it['area']}")
        doc.add_paragraph(f"Характер радіообміну: {it.get('nature','—')}")
        doc.add_paragraph(f"Головний ВЗ: {it.get('main_vz','—')}")
        doc.add_paragraph(f"Підлеглі ВЗ: {it.get('sub_vz','—')}")
        if it.get("period") and it["period"] != "—":
            doc.add_paragraph(f"Період функціонування: {it['period']}")
        doc.add_paragraph()


# ---------- основний сценарій ----------
def main():
    li = load_inputs()
    reference_df = li.reference_df.copy()
    intercepts_df = li.intercepts_df.copy()
    freq_book_path = Path(li.freq_path) if hasattr(li, "freq_path") else Path("Frequencies_63.xlsx")

    freq_file = Path(__file__).resolve().parent / "data" / "freq.txt"
    tokens = read_freq_tokens(freq_file)
    freq4_list = tokens_to_freq4(tokens, reference_df)

    normalize_frequency_column(intercepts_df, reference_df, li.masks_df)

    rows, items = [], []
    for f4 in freq4_list:
        net_name = get_network_name_by_freq(f4, reference_df) or "—"
        modulation, nature, main_vz, sub_vz, area, period = read_ref_fields(freq_book_path, f4)
        p_start, p_end = activity_period_for_freq4(intercepts_df, f4)
        mask3 = mask3_from_reference(reference_df, f4)

        rows.append({"freq4": f4, "modulation": modulation, "period_start": p_start, "period_end": p_end})
        items.append({"freq4": f4, "network_name": net_name, "nature": nature, "main_vz": main_vz,
                      "sub_vz": sub_vz, "mask3": mask3, "area": area, "period": period})

    today_display = datetime.now().strftime("%d.%m.%Y")
    out_dir = Path.cwd() / "build"; out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"Ворожі радіомережі в тилу (63 ОМБр {datetime.now().strftime('%Y-%m-%d')}).docx"

    doc = Document()
    add_header(doc, today_display, unit_name="63 омсбр")
    add_overview_table(doc, rows)
    add_network_blocks(doc, items)
    
    saved = safe_save_docx(doc, out_path)  # як і раніше використовуємо безпечне збереження
    # doc.save(out_path) 
    log.info("[OK] Звіт збережено: %s", saved.resolve)


if __name__ == "__main__":
    main()
