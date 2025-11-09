from pathlib import Path
import json
import pandas as pd
from docx import Document

from src.armorkit.data_loader import load_config, load_inputs
from src.armorkit.normalize_freq import normalize_frequency_column
from src.armorkit.docxutils.safe_save import safe_save_docx
from src.armorkit.dates import (
    parse_period_from_filename,   # -> ("DD.MM.YYYY HH:MM", "DD.MM.YYYY HH:MM")
    build_report_filename,        # -> "Звіт РЕР (... - ...).docx"
)

from src.armorkit.docxutils.styles import (
    set_base_styles,
    add_right_text,
    add_center_block,
    add_block_title,
    add_footer_signature,
)

from src.armorkit.docxutils.tables import (
    add_intercepts_table,
)


def _hour_from_timestr(timestr: str) -> int:
    """
    parse_period_from_filename повертає щось типу '26.10.2025 05:30'.
    Нам потрібно дістати годину як int.
    """
    try:
        # '26.10.2025 05:30' -> ['26.10.2025', '05:30'] -> '05'
        hh = timestr.split()[1].split(":")[0]
        return int(hh)
    except Exception:
        return -1


def _need_1217(period_start: str, period_end: str) -> bool:
    """
    Приймаємо рішення про форму 1.2.17.
    Умова: якщо година початку або кінця періоду == 19.
    """
    return _hour_from_timestr(period_start) == 19 or _hour_from_timestr(period_end) == 19


def _collect_freq_set(reference_df: pd.DataFrame, tags: list[str]) -> set[str]:
    """
    За заданим списком тегів (імена в колонці 'Хто' довідника частот)
    будуємо множину частот (значення з колонки 'Частота').

    Якщо tags порожній -> повертаємо порожню множину.
    """
    if not tags:
        return set()

    sub = reference_df[reference_df["Хто"].isin(tags)]
    freqs = (
        sub["Частота"]
        .dropna()
        .astype(str)
        .unique()
        .tolist()
    )
    return set(freqs)


def _filter_intercepts(intercepts_df: pd.DataFrame, freq_set: set[str]) -> pd.DataFrame:
    """
    1. Залишаємо тільки ті перехоплення, у яких 'Частота' ∈ freq_set.
    2. Залишаємо тільки ті, де в 'примітки' є НЕпорожній текст.
       - не NaN
       - не ""
       - не " " (тільки пробіли)
    3. Сортуємо за 'Дата' + 'Час', якщо обидва стовпці існують.

    Повертає ВЖЕ ВІДФІЛЬТРОВАНИЙ датафрейм.
    """
    if not freq_set:
        # повертаємо пустий датафрейм із тими ж колонками
        return intercepts_df.iloc[0:0]

    df = intercepts_df.copy()

    # 1. Фільтр по множині частот
    df["Частота"] = df["Частота"].astype(str)
    df = df[df["Частота"].isin(freq_set)]

    # 2. Фільтр тільки на наявність коментаря
    #    будуємо допоміжну булеву маску "є валідні примітки"
    if "примітки" in df.columns:
        notes_series = df["примітки"].astype(str).str.strip()
        mask_has_notes = notes_series.ne("") & notes_series.ne("nan") & notes_series.ne("None")
        df = df[mask_has_notes]
    else:
        # якщо стовпця "примітки" взагалі немає — тоді немає що показувати
        return df.iloc[0:0]

    # 3. Сортування
    if {"Дата", "Час"}.issubset(df.columns):
        df = df.sort_values(by=["Дата", "Час"], kind="stable")

    return df


def _add_header_block(doc: Document, include_1217: bool, period_start: str, period_end: str):
    """
    Формує шапку документа:
    1. Якщо include_1217 == True -> праворуч "Форма №1.2.17/ОСУВ".
    2. Три центровані рядки:
          Донесення
          за результатами...
          у зоні...
    3. Рядок "(станом на HH:MM DD.MM.YYYY)".
       Беремо кінець періоду і переставляємо місцями час і дату.
       Наприклад:
         period_end = "26.10.2025 07:00"
         виведемо "(станом на 07:00 26.10.2025)".
    """

    if include_1217:
        add_right_text(doc, "Форма №1.2.17/ОСУВ", bold=False, font_size_pt=12)

    add_center_block(
        doc,
        [
            "Донесення",
            "за результатами ведення радіоелектронної розвідки",
            "у зоні відповідальності 3 АК",
        ],
        bold=False,
        font_size_pt=12,
    )

    # побудувати "(станом на HH:MM DD.MM.YYYY)"
    parts = period_end.split()
    if len(parts) >= 2:
        date_part = parts[0]  # DD.MM.YYYY
        time_part = parts[1]  # HH:MM
        line = f"(станом на {time_part} {date_part})"
    else:
        # якщо чомусь формат інший - просто покажемо як є
        line = f"(станом на {period_end})"

    add_center_block(doc, [line], bold=False, font_size_pt=12)


def _add_area_block(
    doc: Document,
    area_title: str,
    tags: list[str],
    intercepts_df: pd.DataFrame,
    reference_df: pd.DataFrame,
):
    """
    Додає один інформаційний блок у документ.

    Правила:
    - Якщо tags порожній -> повністю пропускаємо блок.
    - Частоти збираємо з reference_df по колонці "Хто".
      Якщо частот для цих тегів немає -> блок пропускаємо.
    - Відібрані перехоплення:
        * мають відповідати цим частотам
        * мають НЕпорожні "примітки" (фільтр робиться в _filter_intercepts)
    - Якщо після фільтрів нічого не лишилось -> друкуємо заголовок +
      "У даній смузі інформативних радіоперехоплень не зафіксовано."
    - Якщо лишились рядки -> друкуємо заголовок + таблицю (р\обмін | примітки)
      з рамками.
    """

    # 1. Немає тегів → цей блок взагалі не публікуємо.
    if not tags:
        return

    # 2. Множина частот для цих тегів
    freq_set = _collect_freq_set(reference_df, tags)
    if not freq_set:
        return  # взагалі не друкуємо блок, бо зона по суті не активна

    # 3. Фактичні перехоплення з валідними "примітками"
    block_df = _filter_intercepts(intercepts_df, freq_set)

    # 4. Публікуємо заголовок блоку в будь-якому разі (бо сектор існує)
    add_block_title(doc, area_title)

    # 5. Якщо після фільтрації нема жодного перехоплення з коментарем:
    if block_df.empty:
        doc.add_paragraph("У даній смузі інформативних радіоперехоплень не зафіксовано.")
        doc.add_paragraph("")  # відступ після блоку
        return

    # 6. Є релевантні перехоплення → будуємо таблицю
    rows_data = []
    for _, row in block_df.iterrows():
        rows_data.append(
            (
                row.get("р\\обмін", ""),
                row.get("примітки", ""),
            )
        )

    add_intercepts_table(doc, rows_data)
    doc.add_paragraph("")  # відступ після таблиці


def build_simple_report_docx(config_path: str) -> Path:
    """
    Основна точка входу.
    1. Завантажує дані через load_inputs / load_config.
    2. Нормалізує перехоплення (маски частот -> реальні частоти).
    3. Визначає часовий діапазон з імені XLSX через parse_period_from_filename.
    4. Генерує docx з усіма блоками.
    5. Зберігає файл у build/ (або cfg.paths.output_dir).
    Повертає шлях до збереженого документа.
    """

    # 1. Дані
    cfg = load_config(config_path)
    li = load_inputs(config_path)

    # print(f"DEBUG: Generating simple report docx with {li.masks_df} intercepts")
    # 2. Нормалізувати частоти:
    normalize_frequency_column(li.intercepts_df, li.reference_df, li.masks_df)

    # 3. Часовий діапазон:
    #    parse_period_from_filename повертає два РЯДКИ:
    #    period_start = "DD.MM.YYYY HH:MM"
    #    period_end   = "DD.MM.YYYY HH:MM"
    period_start, period_end = parse_period_from_filename(li.report_path)

    # 4. Чи треба режим 1.2.17:
    include_1217 = _need_1217(period_start, period_end)

    # 5. Готуємо документ:
    doc = Document()
    set_base_styles(doc)

    # 6. Шапка:
    _add_header_block(doc, include_1217, period_start, period_end)

    # 7. Витягуємо блоки з data.json:
    data_json_path = Path(__file__).parent / "data.json"
    with open(data_json_path, "r", encoding="utf-8") as f:
        cfg_blocks = json.load(f)
        # Очікуваний формат:
        # {
        #   "Areas": [
        #     { "Title": "у зоні відповідальності 63 омбр:", "Tags": ["31 мсп ...", ...] },
        #     ...
        #   ]
        # }

    for block in cfg_blocks.get("Areas", []):
        title = block.get("Title", "").strip()
        tags = block.get("Tags", [])

        _add_area_block(
            doc=doc,
            area_title=title,
            tags=tags,
            intercepts_df=li.intercepts_df,
            reference_df=li.reference_df,
        )

    # 8. Підпис командира:
    add_footer_signature(doc)

    # 9. Ім'я вихідного файлу.
    #    Стандартна функція build_report_filename робить:
    #    "Звіт РЕР (ДД.ММ.РРРР_ГГ-ММ - ДД.ММ.РРРР_ГГ-ММ).docx"
    #    Нам треба врахувати 1.2.17 у префіксі.
    prefix = "Звіт РЕР 1.2.17" if include_1217 else "Звіт РЕР"
    filename = build_report_filename(period_start, period_end, prefix=prefix, sufix="docx")

    # 10. Куди зберігаємо:
    out_dir = Path(getattr(cfg.paths, "output_dir", "build"))
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename

    # 11. Безпечне збереження:
    saved_path = safe_save_docx(doc, out_path)

    return saved_path
