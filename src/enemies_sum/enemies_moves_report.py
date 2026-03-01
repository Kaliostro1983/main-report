from pathlib import Path
from datetime import datetime

import re

import logging
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.armorkit.data_loader import load_config, load_inputs
from src.armorkit.docxutils.safe_save import safe_save_docx

from src.armorkit.data_provider.moves_provider import load_moves_df
from src.armorkit.normalize_freq import normalize_numeric_column


def _enrich_moves_with_reference(moves_df: pd.DataFrame, reference_df: pd.DataFrame) -> pd.DataFrame:
    """
    Додає до moves_df поля 'Підрозділ' та 'Група' (з колонки 'Хто' довідника частот).

    Логіка зіставлення:
      1) спочатку по Маска_3
      2) де не знайшли – по Частота
    """

    df = moves_df.copy()

    # гарантуємо числовий формат для join
    # df["Частота/маска"] = df["Частота/маска"].astype(float)
    df = normalize_numeric_column(df, "Частота/маска")

    ref = reference_df.copy()

    # Частоти/маски як float
    if "Маска_3" in ref.columns:
        ref["Маска_3"] = ref["Маска_3"].astype(float, errors="ignore")
    if "Частота" in ref.columns:
        ref["Частота"] = ref["Частота"].astype(float, errors="ignore")

    # Підмножина колонок з довідника
    cols_for_merge = []
    for col in ["Частота", "Маска_3", "Підрозділ", "Хто"]:
        if col in ref.columns:
            cols_for_merge.append(col)

    ref_sub = ref[cols_for_merge].copy()

    # 1) join по Маска_3
    if "Маска_3" in ref_sub.columns:
        merged = df.merge(
            ref_sub[["Маска_3", "Підрозділ", "Хто"]],
            how="left",
            left_on="Частота/маска",
            right_on="Маска_3",
        )
    else:
        merged = df.copy()
        merged["Підрозділ"] = None
        merged["Хто"] = None

    # 2) fallback: для рядків без Підрозділ пробуємо join по Частота
    # 2) fallback: для рядків без Підрозділ пробуємо join по Частота
    if "Частота" in ref_sub.columns:
        mask_missing = merged["Підрозділ"].isna()
        if mask_missing.any():
            # беремо саме ті рядки, що "порожні", з merged (індекси збігаються)
            missing_rows = merged.loc[mask_missing, ["Частота/маска"]].copy()

            # дедуп довідника по Частота, щоб merge не розмножував рядки
            ref_freq = ref_sub[["Частота", "Підрозділ", "Хто"]].drop_duplicates(subset=["Частота"], keep="first")

            fallback = missing_rows.merge(
                ref_freq,
                how="left",
                left_on="Частота/маска",
                right_on="Частота",
            )

            # присвоєння 1:1 (довжини гарантовано однакові)
            merged.loc[mask_missing, "Підрозділ"] = fallback["Підрозділ"].to_numpy()
            merged.loc[mask_missing, "Хто"] = fallback["Хто"].to_numpy()

    # перейменовуємо "Хто" у "Група"
    if "Хто" in merged.columns:
        merged.rename(columns={"Хто": "Група"}, inplace=True)

    # --- ПОПЕРЕДЖЕННЯ ПРО ЗАПИСИ БЕЗ ГРУПИ ---
    if "Група" in merged.columns:
        missing_group_df = merged[merged["Група"].isna()].copy()
        if not missing_group_df.empty:
            freqs_list = (
                missing_group_df["Частота/маска"]
                .dropna()
                .astype(str)
                .unique()
                .tolist()
            )
            logging.warning(
                "Є %d запис(ів) у moves.xlsx, для яких не знайдено групу в довіднику. Частоти/маски: %s",
                len(missing_group_df),
                ", ".join(freqs_list),
            )
    else:
        logging.warning(
            "У довіднику немає колонки 'Хто', тому поле 'Група' не буде заповнене."
        )

    # чистимо, залишаємо тільки потрібні нам колонки
    keep_cols = ["Частота/маска", "Переміщення", "Підрозділ", "Група"]
    existing_keep_cols = [c for c in keep_cols if c in merged.columns]
    enriched = merged[existing_keep_cols].copy()

    # відкидаємо рядки без групи або без тексту переміщення
    if "Група" in enriched.columns:
        enriched = enriched.dropna(subset=["Група"])
    enriched = enriched.dropna(subset=["Переміщення"])

    # сортування: спочатку група, потім частота/маска (якщо є)
    sort_cols = [c for c in ["Група", "Частота/маска"] if c in enriched.columns]
    if sort_cols:
        enriched = enriched.sort_values(by=sort_cols, kind="stable")

    return enriched



def _set_base_font_12(doc: Document):
    """Встановлює базовий розмір шрифту 12 pt для стилю Normal."""
    style = doc.styles["Normal"]
    font = style.font
    font.size = Pt(12)


def _add_header(doc: Document):
    """
    Додає заголовок:
      'Звіт про переміщення ворога'
      'за результатами радіорозвідки ([сьогоднішня дата])'
    """

    today_str = datetime.now().strftime("%d.%m.%Y")

    # 1-й рядок
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("Звіт про переміщення ворога")
    run1.bold = True
    run1.font.size = Pt(14)

    # 2-й рядок
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"за результатами радіорозвідки ({today_str})")
    run2.font.size = Pt(14)

    # порожній рядок після шапки
    doc.add_paragraph("")


import re


def _extract_time_for_sort(text: str):
    """
    Витягує перший час формату HH:MM з рядка.
    Якщо не знайдено — повертає None (такі підуть в кінець).
    """
    if not isinstance(text, str):
        return None

    m = re.search(r"\b([01]\d|2[0-3]):([0-5]\d)\b", text)
    if not m:
        return None

    return m.group(0)


def _add_group_blocks(doc: Document, enriched_df: pd.DataFrame):
    """
    Створює текстові блоки по кожній групі.

    Для кожної групи:
      В зоні функціонування [назва групи] виявлено наступні переміщення:
      • [Переміщення] (р/м:[Частота/маска], [Підрозділ]).
    """

    if enriched_df.empty:
        doc.add_paragraph("За наявними даними переміщень ворога не зафіксовано.")
        return

    groups = enriched_df["Група"].dropna().unique()

    first_block = True
    for group_name in groups:
        group_df = enriched_df[enriched_df["Група"] == group_name].copy()

        if group_df.empty:
            continue

        # один порожній рядок між блоками
        if not first_block:
            doc.add_paragraph("")
        first_block = False

        # заголовок блоку
        header_text = f"В зоні функціонування {group_name} виявлено наступні переміщення:"
        p = doc.add_paragraph()
        run = p.add_run(header_text)
        run.bold = True

        # ---- сортування за часом усередині блоку ----
        group_df["_sort_time"] = group_df["Переміщення"].apply(_extract_time_for_sort)

        group_df["_sort_key"] = group_df["_sort_time"].fillna("99:99")
        group_df = group_df.sort_values(by="_sort_key", kind="stable")

        # ---- булети ----
        for _, row in group_df.iterrows():
            move_text = str(row.get("Переміщення", "")).strip()
            freq_value = row.get("Частота/маска", "")
            unit = str(row.get("Підрозділ", "")).strip()

            freq_str = "" if pd.isna(freq_value) else str(freq_value)
            unit_str = "" if pd.isna(unit) else unit

            line = f"{move_text} (р/м:{freq_str}, {unit_str})."

            doc.add_paragraph(line, style="List Bullet")



def build_enemy_moves_report_docx(config_path: str) -> Path:
    """
    Основна функція створення документа:

    1. Завантажує config та inputs.
    2. Читає moves.xlsx.
    3. Збагачує дані за довідником частот.
    4. Формує документ.
    5. Зберігає у build/ (або cfg.paths.output_dir).
    """

    # 1. Конфіг та вхідні дані
    cfg = load_config(config_path)
    li = load_inputs(config_path)

    # 2. Завантаження moves.xlsx з поточної папки модуля
    # moves_path = Path(__file__).parent / "moves.xlsx"
    # if not moves_path.exists():
    #     raise FileNotFoundError(f"Файл з переміщеннями не знайдено: {moves_path}")

    # moves_df = pd.read_excel(moves_path)
    
    moves_path = Path(__file__).parent / "moves.xlsx"  # дефолт для xlsx-режиму
    # moves_df = load_moves_df(cfg, default_xlsx_path=moves_path)
    moves_df = load_moves_df(cfg, config_path=config_path, default_xlsx_path=moves_path)
    
    
    from src.armorkit.normalize_freq import normalize_numeric_column

    moves_df = load_moves_df(cfg, config_path=config_path, default_xlsx_path=moves_path)

    # 1) нормалізуємо текст (щоб NaN не став "nan")
    moves_df["Переміщення"] = moves_df["Переміщення"].fillna("").astype(str).str.strip()

    # 2) нормалізуємо число універсально (кома/крапка/NBSP)
    normalize_numeric_column(moves_df, "Частота/маска")

    # 3) відсікаємо хвіст
    moves_df = moves_df[
        moves_df["Частота/маска"].notna()
        & (moves_df["Переміщення"] != "")
    ].copy()
    
    

    # перевіряємо мінімально необхідні колонки
    required_cols = {"Частота/маска", "Переміщення"}
    if not required_cols.issubset(moves_df.columns):
        missing = required_cols - set(moves_df.columns)
        raise ValueError(f"У moves.xlsx відсутні необхідні колонки: {missing}")

    # 3. Збагачуємо дані за довідником частот
    reference_df = li.reference_df
    enriched_df = _enrich_moves_with_reference(moves_df, reference_df)

    # 4. Формуємо документ
    doc = Document()
    _set_base_font_12(doc)
    _add_header(doc)
    _add_group_blocks(doc, enriched_df)

    # 5. Формуємо ім'я файлу та шлях збереження
    today_str = datetime.now().strftime("%d.%m.%Y")
    filename = f"Переміщення ворога ({today_str}).docx"

    out_dir = Path(getattr(cfg.paths, "output_dir", "build"))
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename

    saved_path = safe_save_docx(doc, out_path)
    return saved_path
